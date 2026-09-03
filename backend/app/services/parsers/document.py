import pymupdf as fitz  # Modern PyMuPDF import
import docx
from fastapi import UploadFile, HTTPException
import io

MAX_PAGES = 35
MAX_SIZE_BYTES = 15 * 1024 * 1024 # 15MB

async def extract_text_from_file(file: UploadFile) -> str:
    content = await file.read()
    
    if len(content) > MAX_SIZE_BYTES:
        raise HTTPException(status_code=400, detail="File size exceeds maximum allowed size of 15MB.")
    
    filename = (file.filename or "").lower()
    text = ""

    # 1. Plain text/markdown
    if filename.endswith(".txt") or filename.endswith(".md"):
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

    # 2. Portable Document Format (PDF)
    elif filename.endswith(".pdf"):
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            if len(doc) > MAX_PAGES:
                raise HTTPException(status_code=400, detail=f"PDF exceeds {MAX_PAGES} page limit.")
            
            for page in doc:
                text += page.get_text() + "\n"
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed parsing PDF: {str(e)}")

    # 3. Microsoft Word (DOCX)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        try:
            doc_file = io.BytesIO(content)
            doc = docx.Document(doc_file)
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    text += " | ".join(cell.text.strip() for cell in row.cells) + "\n"
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed parsing Word Document: {str(e)}")

    else:
        raise HTTPException(status_code=400, detail="Unsupported file format.")

    clean_text = text.strip()
    if len(clean_text) < 20:
        raise HTTPException(status_code=400, detail="Document doesn't contain enough readable text.")

    return clean_text[:40000]
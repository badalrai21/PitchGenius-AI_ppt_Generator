import os
import re
import json
import httpx
from fastapi import APIRouter, HTTPException, status, File, UploadFile, Form
from pydantic import BaseModel
from typing import Optional, List
from app.core.config import get_settings
from app.db.supabase import get_supabase_admin
from app.services.ai.discovery import ModelDiscoveryService
from app.services.image_search import ImageSearchService
from app.services.ai.prompts import get_prompt_from_db

router = APIRouter()
settings = get_settings()


# ════════════════════════════════════════════════════════════════
# Pydantic Schemas
# ════════════════════════════════════════════════════════════════
class SlideContent(BaseModel):
    layout: str
    title: str
    subtitle: Optional[str] = None
    bullets: Optional[List[str]] = None
    body: Optional[str] = None
    left_column: Optional[dict] = None
    right_column: Optional[dict] = None
    metrics: Optional[List[dict]] = None
    quote: Optional[str] = None
    quote_author: Optional[str] = None
    image_keyword: Optional[str] = None
    image_url: Optional[str] = None
    visual_style: Optional[str] = None


class GeneratePromptRequest(BaseModel):
    prompt: str
    language: str = "English"
    slide_count: int = 8
    art_style: str = "modern"
    user_id: str


class EnhanceSlideRequest(BaseModel):
    slide: SlideContent
    action: str


class ImageGenerationRequest(BaseModel):
    prompt: str
    use_ai_generation: bool = False
    aspect_ratio: Optional[str] = "16:9"


# ════════════════════════════════════════════════════════════════
# Helpers
# ════════════════════════════════════════════════════════════════

async def get_default_theme_from_db() -> dict:
    """★ Fetch default theme from DB templates table instead of hardcoding."""
    try:
        supabase = get_supabase_admin()
        res = supabase.table("templates").select("theme_config, colors").eq("is_active", True).order("sort_order").limit(1).execute()
        if res.data and len(res.data) > 0:
            theme = res.data[0].get("theme_config") or res.data[0].get("colors") or {}
            if theme:
                return theme
    except Exception as e:
        print(f"[Theme] Could not fetch from DB: {e}")
    # Fallback only if DB is completely empty
    return {
        "primary": "#0077B6", "secondary": "#1D1D1F", "accent": "#86868B",
        "bg": "#FFFFFF", "text": "#1D1D1F", "muted": "#86868B",
        "gradient": "linear-gradient(135deg, #0077B6 0%, #1D1D1F 100%)",
    }


def safe_insert_presentation(payload: dict) -> dict:
    supabase = get_supabase_admin()
    try:
        res = supabase.table("presentations").insert(payload).execute()
        if res.data and len(res.data) > 0:
            return res.data[0]
    except Exception as e:
        err_msg = str(e)
        if "PGRST204" in err_msg or "Could not find" in err_msg:
            safe_payload = {
                "user_id": payload.get("user_id"),
                "title": payload.get("title", "Untitled"),
                "slides_data": payload.get("slides_data", []),
                "slide_count": payload.get("slide_count", 0),
                "status": "completed",
                "source_type": payload.get("source_type", "prompt"),
                "custom_theme": payload.get("custom_theme", {})
            }
            for k in ["language", "art_style", "prompt"]:
                if k in payload and k not in err_msg:
                    safe_payload[k] = payload[k]
            res_retry = supabase.table("presentations").insert(safe_payload).execute()
            if res_retry.data and len(res_retry.data) > 0:
                return res_retry.data[0]
        raise e
    raise Exception("Database insert failed.")


def increment_user_quota(user_id: str):
    """★ Atomically increment user's monthly presentation counter."""
    try:
        supabase = get_supabase_admin()
        supabase.rpc("increment_ppt_count", {"user_uuid": user_id}).execute()
        print(f"[Quota] Incremented ppt_count for user {user_id}")
    except Exception as e:
        print(f"[Quota] Could not increment count: {e}")


async def call_groq_dynamic(system_prompt: str, user_prompt: str, max_tokens: int = 4000, temperature: float = 0.7, json_mode: bool = True) -> str:
    api_key = settings.GROQ_API_KEY
    if not api_key or "your" in api_key.lower():
        raise Exception("Groq API key not configured")
    models = await ModelDiscoveryService.get_groq_models()
    if not models:
        raise Exception("No Groq models discovered")
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    last_error = ""
    for model in models[:5]:
        try:
            payload = {
                "model": model,
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                "temperature": temperature,
                "max_tokens": max_tokens
            }
            if json_mode:
                payload["response_format"] = {"type": "json_object"}
            async with httpx.AsyncClient(timeout=90.0) as client:
                res = await client.post(url, headers=headers, json=payload)
                if res.status_code == 200:
                    print(f"[AI] Groq success: {model}")
                    return res.json()["choices"][0]["message"]["content"]
                last_error = f"{model}: {res.status_code}"
        except Exception as e:
            last_error = f"{model}: {str(e)[:100]}"
    raise Exception(f"All Groq models failed. Last: {last_error}")


async def call_gemini_dynamic(system_prompt: str, user_prompt: str, max_tokens: int = 4000, temperature: float = 0.7) -> str:
    api_key = settings.GOOGLE_GEMINI_API_KEY
    if not api_key or "your" in api_key.lower():
        raise Exception("Gemini API key not configured")
    models = await ModelDiscoveryService.get_gemini_models()
    if not models:
        raise Exception("No Gemini models discovered")
    last_error = ""
    for versioned_model in models[:3]:
        try:
            version, model_id = versioned_model.split("/", 1) if "/" in versioned_model else ("v1beta", versioned_model)
            url = f"https://generativelanguage.googleapis.com/{version}/models/{model_id}:generateContent?key={api_key}"
            payload = {
                "contents": [{"parts": [{"text": f"{system_prompt}\n\n{user_prompt}"}]}],
                "generationConfig": {
                    "temperature": temperature,
                    "maxOutputTokens": max_tokens,
                    "responseMimeType": "application/json"
                }
            }
            async with httpx.AsyncClient(timeout=90.0) as client:
                res = await client.post(url, headers={"Content-Type": "application/json"}, json=payload)
                if res.status_code == 200:
                    print(f"[AI] Gemini success: {versioned_model}")
                    return res.json()["candidates"][0]["content"]["parts"][0]["text"]
                last_error = f"{versioned_model}: {res.status_code}"
        except Exception as e:
            last_error = f"{versioned_model}: {str(e)[:100]}"
    raise Exception(f"All Gemini models failed. Last: {last_error}")


async def generate_ai_text(system_prompt: str, user_prompt: str, max_tokens: int = 4000, temperature: float = 0.7, json_mode: bool = True) -> str:
    try:
        return await call_groq_dynamic(system_prompt, user_prompt, max_tokens, temperature, json_mode)
    except Exception as e_groq:
        print(f"[AI] Groq failed: {e_groq}. Trying Gemini...")
        try:
            return await call_gemini_dynamic(system_prompt, user_prompt, max_tokens, temperature)
        except Exception as e_gemini:
            raise Exception(f"All AI failed. Groq: {e_groq} | Gemini: {e_gemini}")


def extract_and_heal_json(text: str) -> List[dict]:
    clean = text.strip()
    clean = re.sub(r"^```(?:json)?\s*", "", clean)
    clean = re.sub(r"\s*```$", "", clean).strip()
    try:
        return normalize_slides(json.loads(clean))
    except Exception:
        pass
    sa, ea, so, eo = clean.find("["), clean.rfind("]"), clean.find("{"), clean.rfind("}")
    if sa != -1 and (so == -1 or sa < so):
        ext = clean[sa:ea+1] if ea != -1 else clean[sa:]
    elif so != -1:
        ext = clean[so:eo+1] if eo != -1 else clean[so:]
    else:
        raise ValueError("No JSON found")
    try:
        return normalize_slides(json.loads(ext))
    except Exception:
        pass
    healed = ext.rstrip(", ")
    if healed.startswith("[") and not healed.endswith("]"):
        lc = healed.rfind("}")
        if lc != -1:
            healed = healed[:lc+1] + "]"
    elif healed.startswith("{") and not healed.endswith("}"):
        lc = max(healed.rfind('"]'), healed.rfind("}"), healed.rfind("]"))
        if lc != -1:
            healed = healed[:lc+1] + "}"
    try:
        return normalize_slides(json.loads(healed))
    except Exception as e:
        raise ValueError(f"Unparseable JSON: {e}")


def normalize_slides(data) -> List[dict]:
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        for k in ["slides", "slides_data", "presentation", "data"]:
            if k in data and isinstance(data[k], list):
                return data[k]
        if "layout" in data:
            return [data]
    raise ValueError("No slides array found")


async def enrich_slides_with_images(slides: List[dict]) -> List[dict]:
    print(f"[Enrich] Fetching images for {len(slides)} slides...")
    keywords = [s.get("image_keyword", s.get("title", "abstract")) for s in slides]
    image_results = await ImageSearchService.batch_get_images(keywords)
    for slide, result in zip(slides, image_results):
        slide["image_url"] = result.get("url", "")
        slide["photographer_name"] = result.get("photographer_name", "")
        slide["photographer_profile"] = result.get("photographer_profile", "")
        slide["image_source"] = result.get("source", "")
    print(f"[Enrich] All images fetched with attribution data")
    return slides


# ════════════════════════════════════════════════════════════════
# LAYOUT CONTRACT (shared between prompt and document modes)
# ════════════════════════════════════════════════════════════════
LAYOUT_CONTRACT = """
LAYOUT-SPECIFIC STRUCTURAL REQUISITES — CRITICAL:
Each slide MUST choose a 'layout' and populate ALL matching keys with RICH, DETAILED content:

1. "title" layout (use for cover, section breaks):
   - Keys: 'title', 'subtitle', 'body' (engaging 2-3 sentence summary paragraph)

2. "bullets" layout (use for key insights, features, lists):
   - Keys: 'title', 'subtitle', 'bullets' (JSON array of 3-5 DETAILED strings with real data)
   - DO NOT put lists in 'body'. Use the 'bullets' array.

3. "two_column" layout (use for comparisons, pros/cons, problem/solution):
   - Keys: 'title',
     'left_column': {"title": "Column Title", "points": ["Detailed point 1", "Detailed point 2"]},
     'right_column': {"title": "Column Title", "points": ["Detailed point 1", "Detailed point 2"]}

4. "metrics" layout (use for data, statistics, KPIs, financials):
   - Keys: 'title', 'subtitle',
     'metrics': [{"value": "$4.2B", "label": "Market Size", "desc": "2024 global TAM"}, ...]
   - Exactly 3 metric objects with REAL or plausible numbers.

5. "quote" layout (use for testimonials, vision, expert opinions):
   - Keys: 'quote' (full quote string), 'quote_author' (name + title)

EVERY slide MUST also include:
- 'image_keyword': 2-3 word visual concept for image search
- 'visual_style': one of gradient | photo | abstract | minimal | dark | colorful
- 'speaker_notes': what the presenter should say
"""


# ════════════════════════════════════════════════════════════════
# ENDPOINTS
# ════════════════════════════════════════════════════════════════

@router.post("/prompt")
async def generate_from_prompt(request: GeneratePromptRequest):
    try:
        # ★ DB-DRIVEN SYSTEM PROMPT (zero hardcoding)
        db_prompt = await get_prompt_from_db("generate_outline")
        if not db_prompt:
            db_prompt = "You are a world-class presentation designer. Create research-backed, visually rich presentations."

        system_prompt = f"""
{db_prompt}

{LAYOUT_CONTRACT}

CRITICAL EXECUTION PARAMETERS:
1. Return EXACTLY {request.slide_count} slide objects in the 'slides' array.
2. Slide 1 MUST use layout "title" with visual_style "photo".
3. Slides 2-N: Mix layouts — bullets (40%), metrics (20%), two_column (20%), quote (10%).
4. Use REAL industry statistics, market sizes, and verified facts. Never invent numbers.
5. Respond in language: {request.language}.
6. Match aesthetic style: {request.art_style}.
7. Do NOT wrap in markdown fences. Return ONLY a single parseable JSON object.

JSON SCHEMA:
{{
  "title": "Presentation Title",
  "subtitle": "Captivating subtitle",
  "slides": [
    {{
      "index": 1,
      "layout": "title",
      "title": "...",
      "subtitle": "...",
      "body": "...",
      "image_keyword": "visual concept",
      "visual_style": "photo",
      "speaker_notes": "..."
    }}
  ]
}}
"""

        user_prompt = (
            f"Generate exactly {request.slide_count} research-backed slides in {request.language}.\n"
            f"Topic: {request.prompt}\n"
            f"Art style: {request.art_style}"
        )

        response_text = await generate_ai_text(system_prompt, user_prompt, 4000, 0.7, json_mode=True)
        slides_data = extract_and_heal_json(response_text)
        slides_data = await enrich_slides_with_images(slides_data)

        # ★ DB-DRIVEN THEME (zero hardcoding)
        default_theme = await get_default_theme_from_db()

        record = safe_insert_presentation({
            "user_id": request.user_id,
            "title": slides_data[0].get("title", "Untitled"),
            "prompt": request.prompt,
            "language": request.language,
            "slide_count": len(slides_data),
            "slides_data": slides_data,
            "status": "completed",
            "source_type": "prompt",
            "art_style": request.art_style,
            "custom_theme": default_theme,
        })

        increment_user_quota(request.user_id)
        return record

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Generation failure: {str(e)}")


@router.post("/document")
async def generate_from_document(
    language: str = Form("English"),
    slide_count: int = Form(8),
    art_style: str = Form("modern"),
    user_id: str = Form(...),
    file: UploadFile = File(...),
    analysis_depth: str = Form("deep"),
    instructions: Optional[str] = Form(None)
):
    try:
        # ═══════════ STEP 1: PARSE DOCUMENT TEXT ═══════════
        content = await file.read()
        text_content = ""

        if file.filename.endswith(".txt") or file.filename.endswith(".md"):
            text_content = content.decode("utf-8", errors="ignore")
        elif file.filename.endswith(".pdf"):
            try:
                import pymupdf as fitz
            except ImportError:
                import fitz
            doc = fitz.open(stream=content, filetype="pdf")
            # ★ DEEP EXTRACTION: Get ALL pages, not just first few
            text_content = "\n\n".join([
                f"--- PAGE {i+1} ---\n{p.get_text()}"
                for i, p in enumerate(doc)
            ])
            doc.close()
        elif file.filename.endswith(".docx"):
            try:
                from docx import Document
                from io import BytesIO
                doc = Document(BytesIO(content))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text_content = "\n\n".join(paragraphs)
                # ★ Also extract tables from DOCX
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            text_content += f"\n{row_text}"
            except Exception:
                text_content = content.decode("utf-8", errors="ignore")
        else:
            text_content = content.decode("utf-8", errors="ignore")

        # Truncate to safe token limit (roughly 40K chars ≈ 10K tokens)
        if len(text_content) > 40000:
            text_content = text_content[:40000] + "\n\n[... document truncated for token limit ...]"

        source_type = "text" if file.filename == "context.txt" else "document"

        # ═══════════ STEP 2: BUILD DEEP ANALYSIS PROMPT ═══════════
        # ★ DB-DRIVEN SYSTEM PROMPT for document analysis
        db_prompt = await get_prompt_from_db("extract_from_document")
        if not db_prompt:
            db_prompt = await get_prompt_from_db("generate_outline")
        if not db_prompt:
            db_prompt = "You are a world-class presentation designer. Analyze documents deeply and create rich presentations."

        system_prompt = f"""
{db_prompt}

{LAYOUT_CONTRACT}

CRITICAL EXECUTION PARAMETERS:
1. Return EXACTLY {slide_count} slide objects in the 'slides' array.
2. Slide 1 MUST use layout "title" with visual_style "photo".
3. Use diverse layouts: bullets, two_column, metrics, quote — NOT just title slides.
4. Respond in language: {language}.
5. Match aesthetic style: {art_style}.
6. Do NOT wrap in markdown fences. Return ONLY a single parseable JSON object.

ANALYSIS DEPTH: {analysis_depth}
{f"ADDITIONAL INSTRUCTIONS: {instructions}" if instructions else ""}

QUALITY REQUIREMENTS:
- Analyze the ENTIRE document content thoroughly, not just section headings.
- Extract REAL data points, statistics, arguments, conclusions, and key insights.
- Every slide MUST have substantive body content or bullet points with real data from the document.
- Do NOT create slides with just a title and empty body. Each slide must be content-rich.
- Use "metrics" layout when the document contains numbers, percentages, or financial data.
- Use "two_column" layout for comparisons, pros/cons, or problem/solution sections.
- Use "quote" layout for notable statements, testimonials, or expert opinions found in the document.

JSON SCHEMA:
{{
  "title": "Presentation Title Derived From Document",
  "subtitle": "Key insight or summary from the document",
  "slides": [
    {{
      "index": 1,
      "layout": "title",
      "title": "...",
      "subtitle": "...",
      "body": "Comprehensive summary paragraph from document",
      "image_keyword": "visual concept",
      "visual_style": "photo",
      "speaker_notes": "..."
    }}
  ]
}}
"""

        user_prompt = f"""
SOURCE DOCUMENT FOR DEEP ANALYSIS:
\"\"\"
{text_content}
\"\"\"

TASK: Analyze the entire document above deeply. Create exactly {slide_count} professional, content-rich slides in {language}. Extract real insights, data, and arguments from the full document — not just headings. Every slide must have detailed substantive content.
"""

        # ═══════════ STEP 3: GENERATE & ENRICH ═══════════
        response_text = await generate_ai_text(system_prompt, user_prompt, 4000, 0.6, json_mode=True)
        slides_data = extract_and_heal_json(response_text)
        slides_data = await enrich_slides_with_images(slides_data)

        # ★ DB-DRIVEN THEME
        default_theme = await get_default_theme_from_db()

        record = safe_insert_presentation({
            "user_id": user_id,
            "title": slides_data[0].get("title", f"From {file.filename}"),
            "prompt": f"Imported: {file.filename}",
            "language": language,
            "slide_count": len(slides_data),
            "slides_data": slides_data,
            "status": "completed",
            "source_type": source_type,
            "art_style": art_style,
            "custom_theme": default_theme,
            "source_content": text_content[:5000],
        })

        increment_user_quota(user_id)
        return record

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document failure: {str(e)}")


@router.post("/enhance-slide")
async def enhance_slide(request: EnhanceSlideRequest):
    try:
        db_prompt = await get_prompt_from_db("rewrite_slide")
        if not db_prompt:
            db_prompt = "Refine slide per instruction. Keep layout. Include image_keyword. Respond raw JSON."

        response_text = await generate_ai_text(
            db_prompt,
            f"Slide:\n{request.slide.model_dump_json()}\nAction: {request.action}",
            1000, 0.5, json_mode=True
        )
        clean = re.sub(r"```(?:json)?\s*", "", response_text.strip())
        clean = re.sub(r"\s*```", "", clean)
        m = re.search(r"({.*})", clean, re.DOTALL)
        if m:
            clean = m.group(1)
        enhanced = json.loads(clean)

        if enhanced.get("image_keyword") and enhanced.get("image_keyword") != request.slide.image_keyword:
            img_result = await ImageSearchService.get_image(enhanced["image_keyword"])
            enhanced["image_url"] = img_result.get("url", "")
            enhanced["photographer_name"] = img_result.get("photographer_name", "")
            enhanced["photographer_profile"] = img_result.get("photographer_profile", "")
            enhanced["image_source"] = img_result.get("source", "")

        return {"slide": enhanced}
    except Exception as e:
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/image")
async def generate_ai_image(request: ImageGenerationRequest):
    """Search or generate a single image with full attribution data."""
    if not request.prompt.strip():
        raise HTTPException(status_code=400, detail="Empty prompt.")

    result = await ImageSearchService.get_image(
        request.prompt,
        use_ai_generation=request.use_ai_generation
    )

    return {
        "url": result.get("url", ""),
        "photographer_name": result.get("photographer_name", ""),
        "photographer_profile": result.get("photographer_profile", ""),
        "image_source": result.get("source", "unsplash"),
        "fallback": result.get("source") == "picsum",
    }


@router.post("/search-images")
async def search_multiple_images(payload: dict):
    """Search for 6 images for MediaPanel grid."""
    query = payload.get("query", "").strip()
    if not query:
        raise HTTPException(status_code=400, detail="Empty query")

    try:
        api_key = settings.UNSPLASH_ACCESS_KEY
        results = []

        if api_key and "your" not in api_key.lower():
            url = "https://api.unsplash.com/search/photos"
            headers = {"Authorization": f"Client-ID {api_key}"}
            params = {
                "query": query,
                "per_page": 6,
                "orientation": "landscape",
                "content_filter": "high",
            }

            async with httpx.AsyncClient(timeout=10.0) as client:
                res = await client.get(url, headers=headers, params=params)
                if res.status_code == 200:
                    data = res.json()
                    for photo in data.get("results", []):
                        user = photo.get("user", {})
                        username = user.get("username", "")
                        results.append({
                            "url": photo.get("urls", {}).get("regular", ""),
                            "thumb": photo.get("urls", {}).get("small", ""),
                            "photographer_name": user.get("name", "Unsplash Photographer"),
                            "photographer_profile": f"https://unsplash.com/@{username}?utm_source=pitchgenius&utm_medium=referral" if username else "",
                            "image_source": "unsplash",
                        })
                        dl_url = photo.get("links", {}).get("download_location")
                        if dl_url:
                            try:
                                await client.get(dl_url, headers=headers)
                            except Exception:
                                pass

        if not results:
            pexels_key = settings.PEXELS_API_KEY
            if pexels_key and "your" not in pexels_key.lower():
                async with httpx.AsyncClient(timeout=10.0) as client:
                    res = await client.get(
                        "https://api.pexels.com/v1/search",
                        headers={"Authorization": pexels_key},
                        params={"query": query, "per_page": 6, "orientation": "landscape"},
                    )
                    if res.status_code == 200:
                        data = res.json()
                        for photo in data.get("photos", []):
                            results.append({
                                "url": photo.get("src", {}).get("large", ""),
                                "thumb": photo.get("src", {}).get("medium", ""),
                                "photographer_name": photo.get("photographer", "Pexels Photographer"),
                                "photographer_profile": photo.get("photographer_url", ""),
                                "image_source": "pexels",
                            })

        if not results:
            for i in range(6):
                results.append({
                    "url": f"https://picsum.photos/seed/{query.replace(' ', '-')}-{i}/800/600",
                    "thumb": f"https://picsum.photos/seed/{query.replace(' ', '-')}-{i}/400/300",
                    "photographer_name": "",
                    "photographer_profile": "",
                    "image_source": "picsum",
                })

        return {"results": results, "count": len(results)}

    except Exception as e:
        print(f"[Search] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Image search failed: {str(e)}")


@router.post("/suggest-image-prompt")
async def suggest_image_prompt(payload: dict):
    try:
        slide = payload.get("slide", {})
        ctx = f"Title: {slide.get('title', '')}\n"
        if slide.get("subtitle"):
            ctx += f"Subtitle: {slide['subtitle']}\n"
        if slide.get("bullets"):
            ctx += f"Points: {', '.join(slide['bullets'])}\n"

        db_prompt = await get_prompt_from_db("image_prompt")
        if not db_prompt:
            db_prompt = "Suggest 2-3 word image search keyword for this slide. Return ONLY keywords."

        r = await generate_ai_text(
            db_prompt,
            f"Slide:\n{ctx}",
            50, 0.7, json_mode=False
        )
        return {"suggested_prompt": r.strip().strip('"').strip("'")}
    except Exception:
        return {"suggested_prompt": "abstract technology"}


@router.get("/models/discover")
async def discover_available_models():
    g = await ModelDiscoveryService.get_groq_models(force_refresh=True)
    m = await ModelDiscoveryService.get_gemini_models(force_refresh=True)
    return {
        "groq": {"total": len(g), "top_10": g[:10]},
        "gemini": {"total": len(m), "top_5": m[:5]},
    }


@router.get("/images/test")
async def test_image_search(q: str = "startup team"):
    unsplash = await ImageSearchService.search_unsplash(q)
    pexels = await ImageSearchService.search_pexels(q)
    return {
        "query": q,
        "unsplash": unsplash or "Not configured or no results",
        "pexels": pexels or "Not configured or no results",
        "unsplash_key_configured": bool(settings.UNSPLASH_ACCESS_KEY and "your" not in settings.UNSPLASH_ACCESS_KEY.lower()),
        "pexels_key_configured": bool(settings.PEXELS_API_KEY and "your" not in settings.PEXELS_API_KEY.lower()),
    }